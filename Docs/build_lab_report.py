"""Build the live-test lab report as a .docx matching the MS1-MS3 report style.

Mirrors Docs/MS3/build_ms3_report.py (Calibri, blue headings, A4, 1 in margins)
so the lab report sits visually next to the milestone reports.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Docs"
OUT_DOCX = OUT_DIR / "Lab_Report_Ayati_Steffgen.docx"


def set_run(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.03

    for name, size, before, after, color in [
        ("Heading 1", 14, 9, 4, "2E74B5"),
        ("Heading 2", 12, 6, 3, "2E74B5"),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    lp = styles["List Paragraph"]
    lp.font.name = "Calibri"
    lp.font.size = Pt(10.2)
    lp.paragraph_format.space_after = Pt(3)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("NILM Project - Lab Report: Live-Test Software Improvements")
    set_run(r, size=16, bold=True)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run("Modeling, Simulation and Automation of Electrical Energy Systems")
    set_run(r2, size=10.5)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(7)
    r3 = p3.add_run(
        "Soheil Ayati (11153003), Marc Steffgen (11149043)   -   "
        "Live test 16.07.2026   -   Report 22.07.2026"
    )
    set_run(r3, size=9.5, color="555555")


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_lead_para(doc, lead, rest):
    """A paragraph whose first phrase is bold (Today. / Limitation. / ...)."""
    p = doc.add_paragraph()
    r = p.add_run(lead)
    set_run(r, bold=True)
    p.add_run(rest)
    return p


def add_numbered_heading(doc, number, title):
    doc.add_heading(f"{number}. {title}", level=1)


POINTS = [
    (
        "A graphical event timeline",
        "Switch events are written to events.csv and shown on the dashboard as a "
        "scrolling text log. The measured power is charted, but the on/off history "
        "of individual devices is not.",
        "During and after a test it is hard to reconstruct what ran when. Reading a "
        "text log line by line does not show overlap between devices, missed edges, "
        "or a claim that flickered, all of which are obvious at a glance on a timeline.",
        "Render a per-device Gantt-style timeline (one lane per device, coloured on "
        "intervals against a shared time axis) alongside the existing power chart. "
        "The event log already carries the exact edge timestamps and the matched "
        "device, so no new data is needed, only the view. It would also serve as a "
        "post-test artifact for the report and for debugging misclassifications.",
    ),
    (
        "Faster and more precise corrections",
        "When a device's steady state changes shape, the engine can briefly misread "
        "it before a second evidence source corrects it. In the live test at about "
        "16:01 the PV export was released as off for roughly ten seconds, then the "
        "residual monitor re-matched its signature and restored it to on.",
        "A PV generation ramp is a step of delta P with delta Q near zero, which is "
        "indistinguishable in the (P, Q) plane from a small switching supply turning "
        "off, so the off-edge wrongly released the PV claim. The correction is real "
        "but slow, because it waits for the residual monitor's multi-second window to "
        "accumulate before re-matching.",
        "Move the discriminating evidence to the moment of the edge instead of waiting "
        "for the residual monitor. A true switch carries the device's harmonic current "
        "with it; a generation ramp carries none. Applying that harmonic and sign test "
        "at edge time, plus a short hold before a generator claim is released, would "
        "shorten the ten-second correction to a fraction of a second and avoid the "
        "visible flicker, rather than detecting and undoing the mistake after the fact.",
    ),
    (
        "Generic load-class labels for unknown devices",
        "Teaching an unknown load requires the user to type a specific name "
        "(standing_fan_high, coffee_machine_standby, and so on). The signature table "
        "and the identify model are then keyed to that exact family.",
        "In a real installation the operator often does not know the device family, "
        "only that something new is drawing power. Requiring a specific name makes the "
        "system brittle to naming and unhelpful when the device is simply unfamiliar.",
        "Offer an automatic electrical-class label derived from features we already "
        "measure, so an unknown load can be reported as resistive, inductive, "
        "capacitive, or non-linear (switch-mode) without a device name. The sign of Q "
        "separates inductive from capacitive; power factor and THD_I separate a clean "
        "resistive heater from a distorted electronic supply. Our own fingerprint "
        "table shows these classes are cleanly separable on this hardware (THD_I of "
        "about 2 percent for the boiler versus about 170 percent for the laptop). The "
        "user could still refine the auto-label into a specific name later, but the "
        "system would be useful immediately.",
    ),
    (
        "Fully asynchronous, hands-off training",
        "The in-mix teach removed the need to empty the mains, but it still asks the "
        "user for one manual action: switch the unknown device off and back on once so "
        "the settled step gives a model-free measurement of its draw.",
        "Any required user action interrupts an experiment and is impossible for loads "
        "that cannot be toggled (a running fridge, an inverter, a sealed installation). "
        "It also prevents unattended, continuous learning.",
        "Learn passively from the residual history the engine already records, with no "
        "toggle required. When unexplained power appears and stays, capture its settled "
        "level, Q, and harmonic signature over time and teach a provisional signature "
        "from that alone, cross-checking against any natural on/off transitions that "
        "happen to occur. The manual toggle would become an optional way to improve "
        "confidence, not a prerequisite, making training a background process the user "
        "never has to trigger.",
    ),
    (
        "Recognition across a device's operating range",
        "A device is taught at whatever it happened to draw during the capture. A "
        "switching supply or PV inverter whose output varies with load or sunlight can "
        "then fall outside the taught range and stop matching (the laptop taught at "
        "66 W idles at 43 W; PV taught around -13 W was observed at -30 W).",
        "Variable-output loads and generation are exactly the cases NILM finds "
        "hardest, and a single taught operating point does not cover them.",
        "Teach a range rather than a point: record the device across several operating "
        "levels, or let a distinctive harmonic fingerprint taper the match toward idle "
        "(already done for high-THD loads). For PV specifically, the unique and "
        "load-independent discriminator is the sign of power, which no consumer "
        "produces, so generation should be recognised by sign across its whole range "
        "rather than by a taught wattage.",
    ),
]


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("Context", level=1)
    add_para(
        doc,
        "Our live NILM system disaggregates per-appliance power from a single PAC4200 "
        "measurement, logs every switch event with exact timestamps, and learns "
        "unknown devices on the go. It performed well in the live evaluation. This "
        "report collects the concrete weaknesses we observed during those tests and "
        "the improvements we would prioritise next. Each point states what the "
        "software does today, why the current behaviour is a limitation, and the "
        "proposed change.",
    )

    for i, (title, today, limitation, improvement) in enumerate(POINTS, start=1):
        add_numbered_heading(doc, i, title)
        add_lead_para(doc, "Today. ", today)
        add_lead_para(doc, "Limitation. ", limitation)
        add_lead_para(doc, "Improvement. ", improvement)

    add_numbered_heading(doc, 6, "Summary")
    add_para(
        doc,
        "The system met the goals of the exercise. The improvements above target the "
        "gaps we felt during live operation: better visibility (timeline), faster "
        "self-correction (edge-time harmonic evidence), lower operator burden (generic "
        "labels and hands-off learning), and robustness to the hardest loads "
        "(variable-output devices and generation). None require new sensing hardware; "
        "all build on features the pipeline already records.",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
